import os
import zipfile
import tempfile
import io
from datetime import datetime, timedelta
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS

# --- All your document generation libraries ---
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# --- Flask App Setup ---
app = Flask(__name__)
CORS(app)  # Allows the frontend to talk to the backend

# --- Text & Disclaimers ---
DISCLAIMER_TEXT = (
    "This document is a template and is provided for general informational and educational "
    "purposes only. It does not constitute, and is not a substitute for, professional legal "
    "advice. The information contained herein is not legal advice of any kind. The use or "
    "reliance of any information contained in this pack is solely at your own risk. We "
    "encourage you to consult with a qualified legal professional in your jurisdiction to "
    "ensure this agreement meets your specific business needs and complies with all local laws."
)
DISCLAIMER_EN = "This document is a template only and not legal advice. Please consult with a qualified legal professional to ensure compliance with local laws."
DISCLAIMER_MS = "Dokumen ini hanyalah templat dan bukan nasihat undang-undang. Sila berunding dengan profesional undang-undang yang berkelayakan untuk memastikan pematuhan undang-undang tempatan."


# ============================================================================
# ALL 10 OF YOUR DOCUMENT GENERATION FUNCTIONS + HELPER
# (Copied directly from your script)
# ============================================================================

def add_header_footer(doc, company_info):
    section = doc.sections[0]
    
    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = company_info['name']
    header_para.runs[0].font.bold = True
    header_para.runs[0].font.size = Pt(12)
    header_para.runs[0].font.color.rgb = RGBColor(31, 78, 120) # Dark Blue
    
    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = (f"Phone: {company_info['phone']} | Email: {company_info['email']} | "
                   f"Tax ID: {company_info['tax_id']} | Reg: {company_info['reg_no']}")
    
    run = footer_para.add_run(footer_text + " | Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    run = footer_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run = footer_para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    run = footer_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- DOC 1 ---
def create_rental_agreement(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    title = doc.add_heading('EQUIPMENT RENTAL AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ms = doc.add_heading('PERJANJIAN SEWAAN PERALATAN', level=2)
    title_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    disclaimer_para = doc.add_paragraph()
    run = disclaimer_para.add_run(DISCLAIMER_TEXT)
    run.font.size = Pt(8)
    run.font.italic = True
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_heading('1. PARTIES AND EQUIPMENT SCHEDULE', level=1)
    doc.add_heading('1. PIHAK-PIHAK DAN JADUAL PERALATAN', level=2)
    
    doc.add_heading('1.1 Parties to this Agreement', level=3)
    p1 = doc.add_paragraph()
    p1.add_run("This Equipment Rental Agreement is entered into as of <<Date>> between:\n\n")
    p1.add_run(f"PROVIDER: {COMPANY_INFO['name']}\n", 'Strong')
    p1.add_run(f"Registration: {COMPANY_INFO['reg_no']}\n")
    p1.add_run(f"Address: {COMPANY_INFO['address']}\n")
    p1.add_run(f"Phone: {COMPANY_INFO['phone']}\n")
    p1.add_run(f"Email: {COMPANY_INFO['email']}\n\n")
    p1.add_run("CLIENT: <<Client Company Name>>\n", 'Strong')
    p1.add_run("Registration: <<Client Reg No>>\n")
    p1.add_run("Contact: <<Contact Person>>\n")
    p1.add_run("Phone: <<Phone>>\n")
    p1.add_run("Email: <<Email>>")
    
    doc.add_heading('1.1 Pihak-Pihak Kepada Perjanjian Ini', level=3)
    p2 = doc.add_paragraph()
    p2.add_run("Perjanjian ini dimulai pada <<Tarikh>> antara:\n\n")
    p2.add_run(f"PEMBEKAL: {COMPANY_INFO['name']}\n", 'Strong')
    p2.add_run(f"Pendaftaran: {COMPANY_INFO['reg_no']}\n")
    p2.add_run(f"Alamat: {COMPANY_INFO['address']}\n")
    p2.add_run(f"Telefon: {COMPANY_INFO['phone']}\n\n")
    p2.add_run("PELANGGAN: <<Nama Syarikat Pelanggan>>\n", 'Strong')
    p2.add_run("Pendaftaran: <<No. Pendaftaran Pelanggan>>")
    
    doc.add_heading('1.2 Equipment Schedule', level=3)
    doc.add_paragraph("The following equipment ('Equipment') is subject to the terms of this Agreement:")
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    headers = ['Equipment ID', 'Description', 'Make & Model', 'Condition', 'Replacement Value (RM)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    sample_data = [
        ['EQ-001', 'Hydraulic Excavator', 'Caterpillar 320', 'Good, 1,200 hrs', '450000'],
        ['EQ-002', 'Wheel Loader', 'Komatsu WA470', 'Good, 2,100 hrs', '380000'],
        ['EQ-003', 'Air Compressor', 'Atlas Copco', 'Good, 4,500 hrs', '95000']
    ]
    for row_idx, data in enumerate(sample_data, 1):
        for col_idx, value in enumerate(data):
            table.rows[row_idx].cells[col_idx].text = str(value)
    doc.add_paragraph('Total Replacement Value: RM 925,000')
    doc.add_paragraph()
    
    doc.add_heading('2. FINANCIAL TERMS AND SECURITY DEPOSIT', level=1)
    doc.add_heading('2. TERMA KEWANGAN DAN DEPOSIT KESELAMATAN', level=2)
    doc.add_heading('2.1 Detailed Payment Clause', level=3)
    p3 = doc.add_paragraph()
    p3.add_run("Rental Rate / Kadar Sewaan:\n").bold = True
    p3.add_run("- Daily Rate / Kadar Harian: RM <<Amount>> per day\n")
    p3.add_run("- Weekly Rate / Kadar Mingguan: RM <<Amount>> per week\n")
    p3.add_run("- Monthly Rate / Kadar Bulanan: RM <<Amount>> per month\n\n")
    p3.add_run("Payment Schedule / Jadual Pembayaran:\n").bold = True
    p3.add_run("- 50% deposit upon booking confirmation\n")
    p3.add_run("- 50% upon job completion and return of equipment\n\n")
    p3.add_run("Bank Details / Butiran Bank:\n").bold = True
    p3.add_run(f"- Bank: {COMPANY_INFO['bank_name']}\n")
    p3.add_run(f"- Account: {COMPANY_INFO['bank_account']}\n")
    p3.add_run(f"- SWIFT: {COMPANY_INFO['swift_code']}")
    
    doc.add_heading('2.2 Late Payment Penalties', level=3)
    p4 = doc.add_paragraph()
    p4.add_run("Payments not received by the due date shall incur a late fee of 1.5% per month (18% per annum) or the maximum rate permitted by law, whichever is lower. ")
    p4.add_run("Bayaran yang tidak diterima akan dikenakan yuran lewat sebanyak 1.5% setiap bulan.")
    
    doc.add_heading('2.3 Security Deposit', level=3)
    p5 = doc.add_paragraph()
    p5.add_run("The Client shall provide a Security Deposit of RM <<Amount>> prior to delivery. ")
    p5.add_run("This deposit shall be held as security against any damage, loss, theft, or unpaid rent. ")
    p5.add_run("The Security Deposit shall be returned within 14 days of equipment's safe return, ")
    p5.add_run("less any deductions for repairs or outstanding amounts.\n\n")
    p5.add_run("Pelanggan harus memberikan Deposit Keselamatan sebanyak RM <<Jumlah>> sebelum penghantaran. ")
    p5.add_run("Deposit ini akan disimpan sebagai keamanan terhadap sebarang kerosakan, kehilangan, atau kecurian. ")
    p5.add_run("Deposit akan dikembalikan dalam 14 hari selepas pulangan peralatan, tolak sebarang potongan.")
    
    doc.add_heading('3. RISK, LIABILITY AND INSURANCE', level=1)
    doc.add_heading('3. RISIKO, LIABILITI DAN INSURANS', level=2)
    doc.add_heading('3.1 Loss or Damage - Risk Transfer', level=3)
    p6 = doc.add_paragraph()
    p6.add_run("The Client assumes all risk and responsibility for the equipment from the moment of delivery until its return. ")
    p6.add_run("The Client is responsible for the full cost of repair or replacement, valued at the amount specified in the Equipment Schedule, if ")
    p6.add_run("the equipment is lost, stolen, or damaged beyond reasonable wear and tear, regardless of cause.\n\n")
    p6.add_run("Pelanggan memikul semua risiko dan tanggungjawab untuk peralatan sejak penghantaran hingga pulangan. ")
    p6.add_run("Pelanggan bertanggungjawab untuk kos penuh pembaikan atau penggantian jika peralatan hilang atau rosak.")
    
    doc.add_heading('3.2 Mandatory Client Insurance', level=3)
    p7 = doc.add_paragraph()
    p7.add_run("The Client shall procure and maintain at its own expense:\n\n")
    p7.add_run("(i) All-Risk Physical Damage Insurance\n").bold = True
    p7.add_run("- Coverage Amount: Full Replacement Value of all equipment\n")
    p7.add_run("- Named Loss Payee: " + COMPANY_INFO['name'] + "\n")
    p7.add_run("- Deductible: Not to exceed RM 10,000 per occurrence\n\n")
    p7.add_run("(ii) Commercial General Liability Insurance\n").bold = True
    p7.add_run("- Minimum Coverage: RM 2,000,000 per occurrence\n")
    p7.add_run("- Named Additional Insured: " + COMPANY_INFO['name'] + "\n\n")
    p7.add_run("A valid Certificate of Insurance must be provided to the Provider 3 business days prior to delivery.")
    
    doc.add_heading('3.3 Indemnification', level=3)
    p8 = doc.add_paragraph()
    p8.add_run("The Client agrees to indemnify, defend, and hold harmless the Provider, its officers, and employees from and against any and all claims, ")
    p8.add_run("liabilities, damages, losses, costs, and expenses (including reasonable attorney's fees) arising from the ")
    p8.add_run("Client's possession, use, operation, or transportation of the Equipment, except to the extent directly resulting from the ")
    p8.add_run("Provider's gross negligence or willful misconduct.")
    
    doc.add_heading('4. OPERATIONAL RESPONSIBILITIES', level=1)
    doc.add_heading('4. TANGGUNGJAWAB OPERASIONAL', level=2)
    doc.add_heading('4.1 Provider Responsibilities', level=3)
    doc.add_paragraph("The Provider shall:", style='List Bullet')
    doc.add_paragraph("Deliver equipment in good working order", style='List Bullet 2')
    doc.add_paragraph("Ensure routine maintenance prior to delivery", style='List Bullet 2')
    doc.add_paragraph("Provide clear operational instructions", style='List Bullet 2')
    doc.add_paragraph(f"Provide 24/7 support for breakdowns at: {COMPANY_INFO['phone']}", style='List Bullet 2')
    doc.add_heading('4.2 Client Responsibilities', level=3)
    doc.add_paragraph("The Client shall:", style='List Bullet')
    doc.add_paragraph("Perform routine daily maintenance (check fluids, tire pressure, etc.)", style='List Bullet 2')
    doc.add_paragraph("Immediately cease operation and report defects to the Provider", style='List Bullet 2')
    doc.add_paragraph("Ensure only trained and certified personnel operate equipment", style='List Bullet 2')
    doc.add_paragraph("Comply with all safety guidelines and local laws", style='List Bullet 2')
    doc.add_paragraph("Not attempt any repairs without Provider's written consent", style='List Bullet 2')
    doc.add_paragraph("Return equipment in same condition as received, less fair wear and tear", style='List Bullet 2')
    
    doc.add_heading('5. TERMINATION AND RECALL', level=1)
    doc.add_heading('5. PENAMATAN DAN PENGINGAT KEMBALI', level=2)
    doc.add_heading('5.1 Termination for Breach (Default)', level=3)
    doc.add_paragraph("The Provider may terminate immediately and repossess the equipment if:", style='List Bullet')
    doc.add_paragraph("Client fails to pay within 7 days of due date", style='List Bullet 2')
    doc.add_paragraph("Client breaches any material term of this Agreement", style='List Bullet 2')
    doc.add_paragraph("Client uses equipment illegally or in an unsafe manner", style='List Bullet 2')
    doc.add_paragraph("Client's insurance coverage lapses", style='List Bullet 2')
    doc.add_paragraph("Client becomes insolvent or bankrupt", style='List Bullet 2')
    
    doc.add_heading('5.2 Recall Notice', level=3)
    p11a = doc.add_paragraph()
    p11a.add_run("Provider reserves the right to recall any or all equipment upon 10 days written notice to the Client. ")
    p11a.add_run("Pembekal berhak memanggil balik peralatan dengan notis bertulis 10 hari.")
    
    doc.add_heading('5.3 Equipment Return', level=3)
    p12 = doc.add_paragraph()
    p12.add_run("Upon end of rental, Client shall return equipment in same condition as received. ")
    p12.add_run("Provider will conduct final inspection within 5 business days. ")
    p12.add_run("Repair costs will be deducted from Security Deposit. Balance returned within 14 days.")
    
    doc.add_heading('6. GENERAL PROVISIONS', level=1)
    doc.add_heading('6.1 Governing Law', level=3)
    p13 = doc.add_paragraph()
    p13.add_run("This Agreement shall be governed by and construed in accordance with the laws of Malaysia. ")
    p13.add_run("Both parties submit to the exclusive jurisdiction of Malaysian courts.\n\n")
    p13.add_run("Perjanjian ini akan ditadbir oleh undang-undang Malaysia. Kedua-dua pihak menyerah kepada ")
    p13.add_run("bidang kuasa eksklusif mahkamah Malaysia.")
    
    doc.add_page_break()
    doc.add_heading('SIGNATURES / TANDATANGAN', level=1)
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    sig_table = doc.add_table(rows=7, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.rows[0].cells[0].text = 'PROVIDER / PEMBEKAL'
    sig_table.rows[0].cells[1].text = 'CLIENT / PELANGGAN'
    sig_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    sig_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    sig_table.rows[1].cells[0].text = COMPANY_INFO['name']
    sig_table.rows[1].cells[1].text = '<<Client Company Name>>'
    sig_table.rows[2].cells[0].text = '\n\nSigned: _____________________'
    sig_table.rows[2].cells[1].text = '\n\nSigned: _____________________'
    sig_table.rows[3].cells[0].text = 'Name: _____________________'
    sig_table.rows[3].cells[1].text = 'Name: _____________________'
    sig_table.rows[4].cells[0].text = 'Title: _____________________'
    sig_table.rows[4].cells[1].text = 'Title: _____________________'
    sig_table.rows[5].cells[0].text = 'Date: _____________________'
    sig_table.rows[5].cells[1].text = 'Date: _____________________'
    sig_table.rows[6].cells[0].text = '\n\nStamp: [Company Stamp]'
    sig_table.rows[6].cells[1].text = '\n\nStamp: [Company Stamp]'
    
    return doc

# --- DOC 2 ---
def create_booking_form(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    title = doc.add_heading('MACHINERY BOOKING FORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ms = doc.add_heading('BORANG TEMPAHAN PERALATAN', level=2)
    title_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(DISCLAIMER_EN).runs[0].font.size = Pt(8)
    doc.add_paragraph(DISCLAIMER_MS).runs[0].font.size = Pt(8)
    doc.add_paragraph()
    
    table = doc.add_table(rows=31, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)
    
    fields = [
        ('Booking Date', datetime.now().strftime('%d-%m-%Y')),
        ('Booking Reference', '<<BK-YYMMDD-XXXXX>>'),
        ('', ''),
        ('CLIENT COMPANY DETAILS', ''),
        ('Company Name', '<<Company Name>>'),
        ('Registration No.', '<<Registration>>'),
        ('Address', '<<Address>>'),
        ('Phone', '<<Phone>>'),
        ('Email', '<<Email>>'),
        ('', ''),
        ('SITE CONTACT PERSON', ''),
        ('Name', '<<Contact Name>>'),
        ('Phone', '<<Contact Phone>>'),
        ('', ''),
        ('RENTAL SITE', ''),
        ('Site Name', '<<Site Name>>'),
        ('Site Address', '<<Site Address>>'),
        ('', ''),
        ('RENTAL PERIOD & PAYMENT', ''),
        ('Start Date & Time', '<<Date & Time>>'),
        ('End Date & Time', '<<Date & Time>>'),
        ('Duration', '<<Days/Weeks>>'),
        ('Purchase Order (PO) #', '<<PO Number>>'),
        ('', ''),
        ('EQUIPMENT REQUIRED', ''),
        ('Equipment 1', '<<Description>> - Qty: <<Qty>>'),
        ('Equipment 2', '<<Description>> - Qty: <<Qty>>'),
        ('Operator Required?', '☐ Yes  ☐ No'),
        ('', ''),
        ('Authorized By', '<<Name>>'),
        ('Signature', '\n\n_______________________'),
    ]
    
    for i, (field, value) in enumerate(fields):
        cell1 = table.rows[i].cells[0]
        cell2 = table.rows[i].cells[1]
        
        if value == '':
            cell1.merge(cell2)
            cell1.text = field
            cell1.paragraphs[0].runs[0].font.bold = True
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 120)
        else:
            cell1.text = field
            cell2.text = value
    
    return doc

# --- DOC 3 ---
def create_invoice_xlsx(COMPANY_INFO):
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    title_font = Font(size=16, bold=True, color="1F4E78")
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    ws['A1'] = "PROFESSIONAL INVOICE / INVOIS PROFESIONAL"
    ws['A1'].font = title_font
    ws.merge_cells('A1:D1')
    
    ws['A3'] = COMPANY_INFO['name']
    ws['A3'].font = Font(bold=True)
    ws['A4'] = f"Tax ID: {COMPANY_INFO['tax_id']} | Reg: {COMPANY_INFO['reg_no']}"
    ws['A5'] = f"Phone: {COMPANY_INFO['phone']} | Email: {COMPANY_INFO['email']}"
    ws['A6'] = COMPANY_INFO['address']
    
    ws['C3'] = "Invoice Number / No. Invois:"
    ws['D3'] = "<<INV-2025-XXXXX>>"
    ws['C4'] = "Invoice Date / Tarikh Invois:"
    ws['D4'] = datetime.now().strftime("%d-%m-%Y")
    ws['C5'] = "Due Date / Tarikh Luput:"
    ws['D5'] = (datetime.now() + timedelta(days=30)).strftime("%d-%m-%Y")
    
    for cell in ['C3', 'C4', 'C5']:
        ws[cell].alignment = Alignment(horizontal='right')
        ws[cell].font = Font(bold=True)

    ws['A8'] = "Bill To / Bil Kepada:"
    ws['A8'].font = Font(bold=True)
    ws['A9'] = "Client: <<Client Name>>"
    ws['A10'] = "Address: <<Address>>"
    ws['A11'] = "Attn: <<Contact Person>>"
    ws['A12'] = "Reference / Rujukan: <<PO Number>>"
    
    ws['A14'] = "Description / Keterangan"
    ws['B14'] = "Qty / Kuantiti"
    ws['C14'] = "Unit Price / Harga Unit (RM)"
    ws['D14'] = "Amount / Jumlah (RM)"
    
    for cell in ['A14', 'B14', 'C14', 'D14']:
        ws[cell].fill = header_fill
        ws[cell].font = header_font
        ws[cell].border = border
    
    sample_data = [
        ('Equipment Rental - 5 days', 5, 8500),
        ('Operator Service - 5 days', 5, 2000),
        ('Delivery & Pickup', 1, 1500),
    ]
    
    start_row = 15
    for i, (desc, qty, price) in enumerate(sample_data):
        row = start_row + i
        ws[f'A{row}'] = desc
        ws[f'B{row}'] = qty
        ws[f'C{row}'] = price
        ws[f'D{row}'] = f"=B{row}*C{row}"
        ws[f'C{row}'].number_format = '#,##0.00'
        ws[f'D{row}'].number_format = '#,##0.00'
        for col in ['A', 'B', 'C', 'D']:
            ws[f'{col}{row}'].border = border

    end_row = start_row + len(sample_data) - 1
    
    subtotal_row = end_row + 2
    ws[f'C{subtotal_row}'] = "Subtotal / Jumlah Kecil:"
    ws[f'D{subtotal_row}'] = f"=SUM(D{start_row}:D{end_row})"
    ws[f'C{subtotal_row}'].font = Font(bold=True)
    ws[f'D{subtotal_row}'].number_format = '#,##0.00'

    sst_row = subtotal_row + 1
    ws[f'C{sst_row}'] = "SST 6% / Cukai SST 6%:"
    ws[f'D{sst_row}'] = f"=D{subtotal_row}*0.06"
    ws[f'C{sst_row}'].font = Font(bold=True)
    ws[f'D{sst_row}'].number_format = '#,##0.00'

    total_row = sst_row + 1
    ws[f'C{total_row}'] = "TOTAL / JUMLAH (RM):"
    ws[f'D{total_row}'] = f"=D{subtotal_row}+D{sst_row}"
    ws[f'C{total_row}'].font = header_font
    ws[f'C{total_row}'].fill = header_fill
    ws[f'D{total_row}'].font = header_font
    ws[f'D{total_row}'].fill = header_fill
    ws[f'D{total_row}'].number_format = '#,##0.00'
    
    payment_row = total_row + 2
    ws[f'A{payment_row}'] = "Payment Details / Butiran Bayaran:"
    ws[f'A{payment_row}'].font = Font(bold=True)
    ws[f'A{payment_row+1}'] = f"Payment Terms / Terma Bayaran: Net 30 days from invoice date"
    ws[f'A{payment_row+2}'] = f"Bank: {COMPANY_INFO['bank_name']}"
    ws[f'A{payment_row+3}'] = f"Account: {COMPANY_INFO['bank_account']}"
    ws[f'A{payment_row+4}'] = "Thank you for your business! / Terima kasih atas urus niaga anda!"
    ws[f'A{payment_row+4}'].font = Font(italic=True)
    
    return wb

# --- DOC 4 ---
def create_service_log(COMPANY_INFO):
    wb = Workbook()
    ws = wb.active
    ws.title = "Service Log"
    
    columns = [
        'Equipment ID / ID Peralatan',
        'Location / Lokasi',
        'Date / Tarikh',
        'Hour Meter / Bacaan Meter Jam',
        'Service Type / Jenis Servis',
        'Description / Keterangan',
        'Parts Used / Alat Ganti',
        'Labour Cost / Kos Buruh (RM)',
        'Parts Cost / Kos Alat Ganti (RM)',
        'Total Cost / Jumlah Kos (RM)',
        'Next Service Due / Servis Seterusnya'
    ]
    col_widths = [15, 15, 12, 12, 15, 30, 25, 15, 15, 15, 18]
    
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = width
    
    ws['A1'] = "EQUIPMENT SERVICE & MAINTENANCE LOG / LOG SERVIS & PENYELENGGARAAN PERALATAN"
    ws['A1'].font = Font(size=16, bold=True, color="1F4E78")
    ws.merge_cells(f'A1:{chr(64 + len(columns))}1')
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws['A2'] = COMPANY_INFO['name']
    ws.merge_cells(f'A2:{chr(64 + len(columns))}2')
    ws['A2'].alignment = Alignment(horizontal='center')
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=10)
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    for col_num, col_name in enumerate(columns, 1):
        cell = ws.cell(row=4, column=col_num)
        cell.value = col_name
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
    
    sample_data = [
        ('EQ-001', 'Site A', '2025-01-15', 1200, 'Routine', '500-hour service, oil change', 'Oil, Filters', 500, 850, '2025-02-15'),
        ('EQ-002', 'Yard', '2025-01-18', 2100, 'Repair', 'Replaced hydraulic hose', 'Hose #H123', 1200, 450, '2025-01-25'),
        ('EQ-003', 'Site B', '2025-01-20', 4500, 'Inspection', 'Pre-rental safety check', 'None', 300, 0, '2025-03-20'),
    ]
    
    for row_num, row_data in enumerate(sample_data, 5):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border
            if col_num in [8, 9, 10]:
                cell.number_format = '#,##0.00'
        
        total_cell = ws.cell(row=row_num, column=10)
        total_cell.value = f"=H{row_num}+I{row_num}"
    
    return wb

# --- DOC 5 ---
def create_payment_reminder(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}").font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph("<<Client Company Name>>")
    doc.add_paragraph("<<Client Address>>")
    doc.add_paragraph("Attn: <<Contact Person / Accounts Dept>>")
    doc.add_paragraph()
    
    title = doc.add_heading('PAYMENT REMINDER / PERINGATAN PEMBAYARAN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p_subject = doc.add_paragraph()
    p_subject.add_run("Subject: Overdue Payment for Invoice <<INV-XXXXX>>").bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Dear <<Client Name>>,\n\n")
    p.add_run("This is a friendly reminder that payment for the following invoice is now overdue:\n\n")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = "Invoice Number"
    table.rows[0].cells[1].text = "<<INV-XXXXX>>"
    table.rows[1].cells[0].text = "Invoice Date"
    table.rows[1].cells[1].text = "<<Date>>"
    table.rows[2].cells[0].text = "Due Date"
    table.rows[2].cells[1].text = "<<Date>>"
    table.rows[3].cells[0].text = "Amount Due (RM)"
    table.rows[3].cells[1].text = "<<Amount>>"
    table.rows[4].cells[0].text = "Days Overdue"
    table.rows[4].cells[1].text = "<<XX>> days"
    
    for cell in table.rows[3].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("As per our terms, a late payment fee of 1.5% per month may be applied to overdue accounts.\n\n")
    p2.add_run("PLEASE ARRANGE PAYMENT IMMEDIATELY / SILA ATUR PEMBAYARAN DENGAN SERTA-MERTA\n\n").bold = True
    p2.add_run("Payment can be made to:\n")
    p2.add_run(f"Bank: {COMPANY_INFO['bank_name']}\n").bold = True
    p2.add_run(f"Account: {COMPANY_INFO['bank_account']}\n").bold = True
    p2.add_run(f"Account Name: {COMPANY_INFO['name']}\n").bold = True
    
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("Please disregard this notice if payment has already been made. If you have any queries, please contact us at ")
    p3.add_run(f"{COMPANY_INFO['phone']} or {COMPANY_INFO['email']}.\n\n")
    p3.add_run("Thank you,\n\n")
    p3.add_run(f"Accounts Department\n")
    p3.add_run(f"{COMPANY_INFO['name']}")
    
    return doc

# --- DOC 6 ---
def create_customer_portal_form(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    title = doc.add_heading('CUSTOMER PORTAL ACCESS REQUEST FORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ms = doc.add_heading('BORANG PERMOHONAN AKSES PORTAL PELANGGAN', level=2)
    title_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph("Please complete this form to register for online access to your account, where you can track equipment, view invoices, and make payments.")
    doc.add_paragraph()

    table = doc.add_table(rows=12, cols=2)
    table.style = 'Light Grid Accent 1'
    
    fields = [
        ('Company Name', '<<Company Name>>'),
        ('Registration No.', '<<Reg No>>'),
        ('', ''),
        ('PRIMARY AUTHORIZED USER', ''),
        ('Contact Name', '<<Contact Name>>'),
        ('Designation', '<<Job Title>>'),
        ('Email Address (Username)', '<<Email>>'),
        ('Phone Number', '<<Phone>>'),
        ('', ''),
        ('PORTAL ACCESS REQUIRED', ''),
        ('Access Features', '☐ View/Download Invoices\n☐ Make Online Payments\n☐ Track Equipment On-Site\n☐ Log Service Requests'),
        ('Notification Preferences', '☐ Email Notifications ☐ SMS Alerts'),
    ]
    
    for i, (field, value) in enumerate(fields):
        cell1 = table.rows[i].cells[0]
        cell2 = table.rows[i].cells[1]
        
        if value == '':
            cell1.merge(cell2)
            cell1.text = field
            cell1.paragraphs[0].runs[0].font.bold = True
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 120)
        else:
            cell1.text = field
            cell2.text = value
    
    doc.add_paragraph()
    doc.add_heading('Declaration / Pengisytiharan', level=3)
    doc.add_paragraph("I, the undersigned, confirm that I am an authorized representative of the above-named company and request access to the customer portal.")
    
    doc.add_paragraph("\n\nSignature: _____________________")
    doc.add_paragraph("Name: _____________________")
    doc.add_paragraph("Date: _____________________")
    doc.add_paragraph("Company Stamp: ")
    
    return doc

# --- DOC 7 ---
def create_quotation_template(COMPANY_INFO):
    wb = Workbook()
    ws = wb.active
    ws.title = "Quotation"
    
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    title_font = Font(size=16, bold=True, color="1F4E78")
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    ws['A1'] = "QUOTATION / SEBUT HARGA"
    ws['A1'].font = title_font
    ws.merge_cells('A1:D1')
    
    ws['A3'] = COMPANY_INFO['name']
    ws['A3'].font = Font(bold=True)
    ws['A4'] = f"Tax ID: {COMPANY_INFO['tax_id']} | Reg: {COMPANY_INFO['reg_no']}"
    ws['A5'] = f"Phone: {COMPANY_INFO['phone']} | Email: {COMPANY_INFO['email']}"
    ws['A6'] = COMPANY_INFO['address']
    
    ws['C3'] = "Quote #:"
    ws['D3'] = "<<QT-YYMMDD-XXXXX>>"
    ws['C4'] = "Quote Date:"
    ws['D4'] = datetime.now().strftime("%d-%m-%Y")
    ws['C5'] = "Valid Until:"
    ws['D5'] = (datetime.now() + timedelta(days=30)).strftime("%d-%m-%Y")
    
    for cell in ['C3', 'C4', 'C5']:
        ws[cell].alignment = Alignment(horizontal='right')
        ws[cell].font = Font(bold=True)

    ws['A8'] = "Client Information:"
    ws['A8'].font = Font(bold=True)
    ws['A9'] = "Client: <<Client Name>>"
    ws['A10'] = "Address: <<Address>>"
    ws['A11'] = "Attn: <<Contact Person>>"
    
    ws['A13'] = "Description"
    ws['B13'] = "Duration"
    ws['C13'] = "Rate (RM)"
    ws['D13'] = "Amount (RM)"
    
    for cell in ['A13', 'B13', 'C13', 'D13']:
        ws[cell].fill = header_fill
        ws[cell].font = header_font
        ws[cell].border = border
    
    sample_data = [
        ('Hydraulic Excavator', '5 days', 8500),
        ('Operator Service', '5 days', 2000),
    ]
    
    start_row = 14
    for i, (desc, duration, rate) in enumerate(sample_data):
        row = start_row + i
        ws[f'A{row}'] = desc
        ws[f'B{row}'] = duration
        ws[f'C{row}'] = rate
        ws[f'D{row}'] = f"={rate}*5"
        ws[f'C{row}'].number_format = '#,##0.00'
        ws[f'D{row}'].number_format = '#,##0.00'
        for col in ['A', 'B', 'C', 'D']:
            ws[f'{col}{row}'].border = border
            
    end_row = start_row + len(sample_data) - 1

    subtotal_row = end_row + 2
    ws[f'C{subtotal_row}'] = "Subtotal:"
    ws[f'D{subtotal_row}'] = f"=SUM(D{start_row}:D{end_row})"
    ws[f'C{subtotal_row}'].font = Font(bold=True)
    ws[f'D{subtotal_row}'].number_format = '#,##0.00'

    sst_row = subtotal_row + 1
    ws[f'C{sst_row}'] = "SST 6%:"
    ws[f'D{sst_row}'] = f"=D{subtotal_row}*0.06"
    ws[f'C{sst_row}'].font = Font(bold=True)
    ws[f'D{sst_row}'].number_format = '#,##0.00'

    total_row = sst_row + 1
    ws[f'C{total_row}'] = "TOTAL QUOTATION (RM):"
    ws[f'D{total_row}'] = f"=D{subtotal_row}+D{sst_row}"
    ws[f'C{total_row}'].font = header_font
    ws[f'C{total_row}'].fill = header_fill
    ws[f'D{total_row}'].font = header_font
    ws[f'D{total_row}'].fill = header_fill
    ws[f'D{total_row}'].number_format = '#,##0.00'
    
    terms_row = total_row + 2
    ws[f'A{terms_row}'] = "Terms & Conditions:"
    ws[f'A{terms_row}'].font = Font(bold=True)
    ws[f'A{terms_row+1}'] = "1. Payment Terms: 50% deposit upon confirmation, 50% upon completion."
    ws[f'A{terms_row+2}'] = "2. Validity: This quotation is valid for 30 days."
    ws[f'A{terms_row+3}'] = "3. Client to provide mandatory 'All-Risk' insurance."
    
    return wb

# --- DOC 8 ---
def create_delivery_checklist(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    title = doc.add_heading('EQUIPMENT PRE-DELIVERY CHECKLIST', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ms = doc.add_heading('SENARAI SEMAK PRA-PENGHANTARAN PERALATAN', level=2)
    title_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    
    table.rows[0].cells[0].text = "Delivery Date"
    table.rows[0].cells[1].text = "<<Date>>"
    table.rows[1].cells[0].text = "Equipment ID"
    table.rows[1].cells[1].text = "<<Equipment ID>>"
    table.rows[2].cells[0].text = "Client"
    table.rows[2].cells[1].text = "<<Client Name>>"
    table.rows[3].cells[0].text = "Delivery Address"
    table.rows[3].cells[1].text = "<<Address>>"
    table.rows[4].cells[0].text = "Inspector"
    table.rows[4].cells[1].text = "<<Inspector Name>>"
    
    doc.add_paragraph()
    doc.add_heading('CHECKLIST / SENARAI SEMAK', level=2)
    
    checklist = [
        'Equipment Cleanliness - Clean and free of debris',
        'Fluid Levels - Oil, water, hydraulic at proper levels',
        'Tires & Wheels - No damage, proper pressure',
        'Engine Start - Starts smoothly, no unusual noise',
        'Brake System - Responsive braking, no leaks',
        'Hydraulic System - No leaks, operates smoothly',
        'Lights & Indicators - All working properly',
        'Safety Features - Guards, mirrors, horn, beacon intact',
        'Hour Meter Reading Recorded - <<Hours>>',
        'Documentation - Manuals and records provided to client',
        'Photos Taken - All angles, interior, and any existing damage',
    ]
    
    for item in checklist:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\n  Status: ☐ OK  ☐ Issue (See notes)")
    
    doc.add_paragraph()
    doc.add_heading('Notes / Nota:', level=3)
    doc.add_paragraph("_________________________________________________________________")
    doc.add_paragraph("_________________________________________________________________")
    doc.add_paragraph()
    
    doc.add_heading('Sign-off / Pengesahan:', level=3)
    doc.add_paragraph("We confirm the equipment listed above has been inspected and delivered in good working order.")
    doc.add_paragraph("Kami mengesahkan peralatan di atas telah diperiksa dan dihantar dalam keadaan baik.")
    
    doc.add_paragraph("\n\nInspector Signature: _____________________")
    doc.add_paragraph("Client Signature: _____________________")
    doc.add_paragraph("Date: _____________________")
    
    return doc

# --- DOC 9 ---
def create_user_guide_pdf(COMPANY_INFO):
    # This function creates a PDF in memory
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor('#1F4E78'), alignment=1, spaceAfter=14)
    h2_style = ParagraphStyle('CustomH2', parent=styles['Heading2'], fontSize=12, textColor=colors.HexColor('#1F4E78'), spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle('CustomBody', parent=styles['BodyText'], fontSize=10, leading=14, spaceAfter=6)
    bullet_style = ParagraphStyle('CustomBullet', parent=body_style, leftIndent=0.25*inch, bulletIndent=0.1*inch, spaceAfter=4)
    disclaimer_style = ParagraphStyle('Disclaimer', parent=body_style, fontSize=9, textColor=colors.darkred, spaceAfter=10)

    story.append(Paragraph(f"<b>{COMPANY_INFO['name']}</b>", styles['Normal']))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph("BILINGUAL BUSINESS OPERATIONS TEMPLATE PACK", title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("<b>Complete Professional Solution for Malaysian Machinery Rental SMEs</b>", h2_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"<b>LEGAL DISCLAIMER:</b> {DISCLAIMER_TEXT}", disclaimer_style))
    story.append(Paragraph("<b>WHAT'S INCLUDED:</b>", h2_style))
    story.append(Paragraph("<b>1. Equipment Rental Agreement</b> - Complete legal contract (EN/MS)", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>2. Machinery Booking Form</b> - Professional bilingual form", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>3. Professional Invoice</b> - Auto-calculated Excel template (EN/MS)", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>4. Equipment Service Log</b> - Maintenance tracking spreadsheet (EN/MS)", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>5. Payment Reminder Letter</b> - Overdue payment follow-up (EN/MS)", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>6. Customer Portal Form</b> - Online client access registration", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>7. Quotation Template</b> - Professional quotes with auto-calc", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>8. Delivery Checklist</b> - Pre-delivery inspection (EN/MS)", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>9. User Guide (This Doc)</b> - Complete instructions", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>10. Product Overview</b> - Sales materials", bullet_style, bulletText='•'))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("<b>HOW TO USE:</b>", h2_style))
    story.append(Paragraph("<b>STEP 1:</b> Open the desired template in Microsoft Word/Excel or Google Docs/Sheets.", bullet_style, bulletText='1.'))
    story.append(Paragraph("<b>STEP 2:</b> Find all text marked with <b>&lt;&lt;placeholders&gt;&gt;</b>.", bullet_style, bulletText='2.'))
    story.append(Paragraph("<b>STEP 3:</b> Replace the placeholders with your client's or job's specific information.", bullet_style, bulletText='3.'))
    story.append(Paragraph("<b>STEP 4:</b> For Excel files (Invoice, Quote, Log), enter your data and the formulas will auto-calculate.", bullet_style, bulletText='4.'))
    story.append(Paragraph("<b>STEP 5:</b> Save the document with a new name (e.g., 'Invoice_ClientName_Date.xlsx').", bullet_style, bulletText='5.'))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("<b>RECOMMENDED WORKFLOW:</b>", h2_style))
    story.append(Paragraph("Client inquires → Use <b>Booking Form</b> to capture details (including PO #).", bullet_style, bulletText='1.'))
    story.append(Paragraph("Confirm booking → Send <b>Quotation Template</b>.", bullet_style, bulletText='2.'))
    story.append(Paragraph("Client confirms → Send <b>Equipment Rental Agreement</b> for signature.", bullet_style, bulletText='3.'))
    story.append(Paragraph("Before delivery → Use <b>Delivery Checklist</b> for inspection.", bullet_style, bulletText='4.'))
    story.append(Paragraph("Job completion → Issue <b>Professional Invoice</b> (use PO # as reference).", bullet_style, bulletText='5.'))
    story.append(Paragraph("Payment overdue → Send <b>Payment Reminder Letter</b>.", bullet_style, bulletText='6.'))
    story.append(Paragraph("After service → Update <b>Equipment Service Log</b> (note Hour Meter).", bullet_style, bulletText='7.'))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("<b>CUSTOMIZATION TIPS:</b>", h2_style))
    story.append(Paragraph("Your main company details (name, address, bank) are already included.", bullet_style, bulletText='•'))
    story.append(Paragraph("You can add your company logo to the headers of the Word documents.", bullet_style, bulletText='•'))
    story.append(Paragraph("Adjust tax rates (e.g., SST 6%) in the Excel formulas if needed.", bullet_style, bulletText='•'))
    story.append(Paragraph("<b>Crucial:</b> Review the legal clauses in the Agreement with a professional.", bullet_style, bulletText='•'))
    
    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    doc_template = SimpleDocTemplate(pdf_buffer, pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    doc_template.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

# --- DOC 10 ---
def create_product_overview(COMPANY_INFO):
    doc = Document()
    add_header_footer(doc, COMPANY_INFO)
    
    title = doc.add_heading('COMPLETE BILINGUAL BUSINESS TEMPLATE PACK', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('FOR MALAYSIAN MACHINERY RENTAL SMEs', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Stop worrying about managing your rental business. Our comprehensive template pack provides everything needed to run professional, compliant, and profitable rental operations.\n\n")
    p.add_run("Berhenti bimbang tentang menguruskan perniagaan sewaan anda. Pakej templat komprehensif kami menyediakan semua yang diperlukan untuk menjalankan operasi sewaan yang profesional dan menguntungkan.\n\n")
    
    doc.add_heading('KEY BENEFITS / FAEDAH UTAMA:', level=2)
    p_benefits = doc.add_paragraph()
    p_benefits.add_run("✓ All-in-One Solution - 10 essential templates\n")
    p_benefits.add_run("✓ Professional Branding - Your company info and logo on every document\n")
    p_benefits.add_run("✓ Bilingual Ready - English + Bahasa Melayu for clear communication\n")
    p_benefits.add_run("✓ Auto-Calculations - Formulas in Excel templates for Invoices & Quotes\n")
    p_benefits.add_run("✓ Legal Protection - A comprehensive rental agreement included\n")
    p_benefits.add_run("✓ Multiple Formats - DOCX, XLSX, PDF\n")
    p_benefits.add_run("✓ Time Saving - Pre-formatted, just fill in client details\n")
    p_benefits.add_run("✓ Cost Effective - Professional quality documents instantly\n")
    
    doc.add_page_break()
    doc.add_heading('PACKAGE CONTENTS / KANDUNGAN PAKEJ:', level=2)
    contents = [
        '01. Equipment Rental Agreement - 12+ page complete contract (EN/MS)',
        '02. Machinery Booking Form - Professional bilingual form',
        '03. Professional Invoice - Auto-calculated with SST (EN/MS)',
        '04. Equipment Service Log - Detailed maintenance tracking (EN/MS)',
        '05. Payment Reminder Letter - Overdue invoice follow-up (EN/MS)',
        '06. Customer Portal Form - Client online access registration',
        '07. Quotation Template - Professional quote system (Excel)',
        '08. Delivery Checklist - Pre-delivery equipment inspection (EN/MS)',
        '09. User Guide - Complete instructions and workflow (PDF)',
        '10. Product Overview (This Doc) - Sales materials and description',
    ]
    for content in contents:
        doc.add_paragraph(content, style='List Bullet')
    
    doc.add_paragraph()
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_final.add_run("Ready to Professionalize Your Operations?\nSedia untuk Memprofesionalkan Operasi Anda?\n\n").bold = True
    p_final.add_run("Get instant access to all 10 professional templates.\n")
    p_final.add_run("Dapatkan akses segera ke semua 10 templat profesional.")
    
    return doc


# ============================================================================
# 5. MAIN FLASK ROUTE
# This is the "backend system" you described
# ============================================================================
@app.route('/generate-pack', methods=['POST'])
def generate_pack():
    try:
        # 1. Get customer details from the form
        COMPANY_INFO = request.json
        
        # 2. Create a temporary directory to store files
        with tempfile.TemporaryDirectory() as temp_dir:
            
            # --- 3. Generate all 10 documents ---
            
            # Doc 1: Agreement
            doc1 = create_rental_agreement(COMPANY_INFO)
            doc1.save(os.path.join(temp_dir, '01_Equipment_Rental_Agreement.docx'))
            
            # Doc 2: Booking Form
            doc2 = create_booking_form(COMPANY_INFO)
            doc2.save(os.path.join(temp_dir, '02_Machinery_Booking_Form.docx'))
            
            # Doc 3: Invoice
            wb3 = create_invoice_xlsx(COMPANY_INFO)
            wb3.save(os.path.join(temp_dir, '03_Professional_Invoice.xlsx'))
            
            # Doc 4: Service Log
            wb4 = create_service_log(COMPANY_INFO)
            wb4.save(os.path.join(temp_dir, '04_Equipment_Service_Log.xlsx'))
            
            # Doc 5: Payment Reminder
            doc5 = create_payment_reminder(COMPANY_INFO)
            doc5.save(os.path.join(temp_dir, '05_Payment_Reminder_Letter.docx'))
            
            # Doc 6: Customer Portal Form
            doc6 = create_customer_portal_form(COMPANY_INFO)
            doc6.save(os.path.join(temp_dir, '06_Customer_Portal_Form.docx'))
            
            # Doc 7: Quotation
            wb7 = create_quotation_template(COMPANY_INFO)
            wb7.save(os.path.join(temp_dir, '07_Quotation_Template.xlsx'))
            
            # Doc 8: Delivery Checklist
            doc8 = create_delivery_checklist(COMPANY_INFO)
            doc8.save(os.path.join(temp_dir, '08_Delivery_Checklist.docx'))
            
            # Doc 9: User Guide (PDF from memory)
            pdf_buffer = create_user_guide_pdf(COMPANY_INFO)
            with open(os.path.join(temp_dir, '09_User_Guide.pdf'), 'wb') as f:
                f.write(pdf_buffer.getvalue())
            
            # Doc 10: Product Overview
            doc10 = create_product_overview(COMPANY_INFO)
            doc10.save(os.path.join(temp_dir, '10_Product_Overview.docx'))
            
            # --- 4. Create the ZIP file ---
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        zf.write(os.path.join(root, file), arcname=file)
            
            zip_buffer.seek(0)
            
            # --- 5. Send the ZIP file to the customer ---
            return send_file(
                zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name='Bilingual_Business_Template_Pack.zip'
            )

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

# ============================================================================
# 6. RUN THE SERVER
# ============================================================================
if __name__ == '__main__':
    app.run(debug=True, port=5000)